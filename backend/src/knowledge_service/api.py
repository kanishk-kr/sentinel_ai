"""
SENTINEL — Knowledge Service API Router
Endpoints for file upload, document classification, RAG queries, and extraction viewing.
"""
from __future__ import annotations

import logging
import shutil
import uuid

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from src.knowledge_service.rag import document_classifier, evidence_resolver, rag_service
from src.shared.auth import get_current_user, require_admin
from src.shared.config import get_settings
from src.shared.database import get_db
from src.shared.models import (
    AccessTagStatus,
    DocumentExtraction,
    KBDocument,
    User,
)
from src.shared.schemas import (
    ClassifyDocumentRequest,
    ExtractionResponse,
    FileDetailResponse,
    FileUploadResponse,
    RAGQueryRequest,
    RAGQueryResponse,
    RAGChunkResult,
)

logger = logging.getLogger(__name__)
settings = get_settings()
router = APIRouter(prefix="/api/v1", tags=["Knowledge Service"])


@router.post("/files/upload", response_model=FileUploadResponse)
async def upload_file(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Upload a document (FR5.1).
    access_tag is NEVER accepted as user-supplied metadata — set to 'unclassified'
    and requires admin review before indexing.
    """
    # Validate file size
    contents = await file.read()
    if len(contents) > settings.max_upload_size_mb * 1024 * 1024:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File too large. Max: {settings.max_upload_size_mb}MB",
        )

    # Determine file type
    filename = file.filename or "unknown"
    file_type = filename.rsplit(".", 1)[-1].lower() if "." in filename else "unknown"

    # Save file
    file_id = str(uuid.uuid4())
    save_path = settings.upload_path / f"{file_id}.{file_type}"
    with open(save_path, "wb") as f:
        f.write(contents)

    # Create DB record — access_tag defaults to 'unclassified', pending admin review
    doc = KBDocument(
        title=filename,
        original_filename=filename,
        source_path=str(save_path),
        file_type=file_type,
        file_size_bytes=len(contents),
        access_tag="unclassified",
        access_tag_status=AccessTagStatus.PENDING_ADMIN_REVIEW,
        processing_status="pending",
    )
    db.add(doc)
    await db.flush()
    await db.refresh(doc)

    return FileUploadResponse(
        id=str(doc.id),
        filename=filename,
        file_type=file_type,
        file_size_bytes=len(contents),
        access_tag=doc.access_tag,
        access_tag_status=doc.access_tag_status.value,
        processing_status=doc.processing_status,
    )


@router.get("/files", response_model=list[FileUploadResponse])
async def list_files(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """List all documents."""
    result = await db.execute(select(KBDocument).order_by(KBDocument.created_at.desc()))
    docs = result.scalars().all()
    return [
        FileUploadResponse(
            id=str(d.id),
            filename=d.filename if hasattr(d, 'filename') else d.title,
            file_type=d.file_type,
            file_size_bytes=d.file_size_bytes,
            access_tag=d.access_tag,
            access_tag_status=d.access_tag_status.value,
            processing_status=d.processing_status,
        )
        for d in docs
    ]


@router.get("/files/{file_id}", response_model=FileDetailResponse)
async def get_file(
    file_id: str,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get file metadata, classification, and extractions."""
    result = await db.execute(select(KBDocument).where(KBDocument.id == uuid.UUID(file_id)))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")

    ext_result = await db.execute(
        select(DocumentExtraction)
        .where(DocumentExtraction.document_id == doc.id)
        .order_by(DocumentExtraction.page_number.asc())
    )
    extractions = ext_result.scalars().all()

    return FileDetailResponse(
        id=str(doc.id),
        title=doc.title,
        original_filename=doc.original_filename,
        file_type=doc.file_type,
        file_size_bytes=doc.file_size_bytes,
        page_count=doc.page_count,
        access_tag=doc.access_tag,
        access_tag_status=doc.access_tag_status.value,
        chunk_count=doc.chunk_count,
        processing_status=doc.processing_status,
        extractions=[
            ExtractionResponse(
                id=str(e.id),
                page_number=e.page_number,
                region_type=e.region_type.value,
                field_name=e.field_name,
                field_value=e.field_value,
                evidence_confidence=e.evidence_confidence_json,
                bbox=e.bbox_json,
                method=e.method.value,
            )
            for e in extractions
        ],
        ingested_at=doc.ingested_at,
    )


@router.post("/admin/documents/{doc_id}/classify")
async def classify_document(
    doc_id: str,
    request: ClassifyDocumentRequest,
    user: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    """
    Admin-only: confirm/override access_tag before indexing (FR5.1).
    This is the ONLY way to set a document's access_tag.
    """
    result = await db.execute(select(KBDocument).where(KBDocument.id == uuid.UUID(doc_id)))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    doc.access_tag = request.access_tag
    doc.access_tag_status = AccessTagStatus.CONFIRMED
    doc.classified_by = user.id
    doc.classification_notes = request.notes
    await db.flush()

    # Now trigger ingestion into vector store
    if doc.processing_status == "pending" or doc.processing_status == "completed":
        try:
            # Read file and create chunks
            from pathlib import Path
            file_path = Path(doc.source_path)
            if file_path.exists():
                text_content = ""
                if doc.file_type == "pdf":
                    try:
                        from PyPDF2 import PdfReader
                        reader = PdfReader(str(file_path))
                        doc.page_count = len(reader.pages)
                        for i, page in enumerate(reader.pages):
                            text_content += f"\n--- Page {i+1} ---\n{page.extract_text() or ''}"
                    except Exception:
                        text_content = "PDF text extraction failed"
                elif doc.file_type in ("txt", "md", "csv"):
                    text_content = file_path.read_text(errors="ignore")
                elif doc.file_type in ("docx",):
                    try:
                        from docx import Document
                        d = Document(str(file_path))
                        text_content = "\n".join(p.text for p in d.paragraphs)
                    except Exception:
                        text_content = "DOCX extraction failed"
                else:
                    text_content = f"File type {doc.file_type} — binary content"

                # Chunk the text
                chunks = _chunk_text(text_content, chunk_size=500, overlap=50)
                chunk_dicts = [
                    {"text": chunk, "title": doc.title, "page_number": i + 1}
                    for i, chunk in enumerate(chunks)
                ]

                # Ingest into vector store
                ingested = await rag_service.ingest_document(
                    document_id=str(doc.id),
                    chunks=chunk_dicts,
                    access_tag=request.access_tag,
                    db=db,
                )
                doc.chunk_count = ingested
                doc.processing_status = "completed"
                await db.flush()

        except Exception as e:
            logger.error(f"Document ingestion failed: {e}")
            doc.processing_status = "failed"
            doc.processing_error = str(e)
            await db.flush()

    return {
        "id": str(doc.id),
        "access_tag": doc.access_tag,
        "access_tag_status": doc.access_tag_status.value,
        "processing_status": doc.processing_status,
        "chunk_count": doc.chunk_count,
    }


@router.post("/rag/query", response_model=RAGQueryResponse)
async def rag_query(
    request: RAGQueryRequest,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Permission-filtered KB query (FR5.2).
    Role from JWT determines access_tags. Filtered BEFORE scoring.
    """
    # Get user's access tags (from JWT / user record)
    user_tags = user.access_tags or ["general"]

    # If admin, can access all
    from src.shared.models import UserRole
    if user.role == UserRole.ADMIN:
        user_tags = ["*"]

    # Additional filter tags from request
    if request.filter_tags:
        user_tags = [t for t in request.filter_tags if t in user_tags or "*" in user_tags]

    # Execute RAG query
    result = await rag_service.query_with_answer(
        query=request.query,
        user_tags=user_tags,
        top_k=request.top_k,
    )

    return RAGQueryResponse(
        query=request.query,
        answer=result.get("answer", ""),
        model_used=result.get("model_used", ""),
        chunks=[
            RAGChunkResult(
                document_id=c.get("document_id", ""),
                document_title=c.get("document_title", ""),
                chunk_text=c.get("chunk_text", ""),
                page_number=c.get("page_number"),
                access_tag=c.get("access_tag", ""),
                similarity_score=c.get("similarity_score", 0.0),
                evidence_confidence=c.get("evidence_confidence"),
            )
            for c in result.get("chunks", [])
        ],
        citations=result.get("citations", []),
        evidence_confidence={
            "overall": "HIGH" if result.get("verified") else "LOW",
            "source_count": len(result.get("chunks", [])),
            "citation_count": len(result.get("citations", [])),
        },
        verified=result.get("verified", False),
    )


def _chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """Split text into overlapping chunks."""
    if not text:
        return []
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk = " ".join(words[i:i + chunk_size])
        if chunk.strip():
            chunks.append(chunk)
        i += chunk_size - overlap
    return chunks
